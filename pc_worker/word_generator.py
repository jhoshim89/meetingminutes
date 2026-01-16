"""
Word Document Generator for Meeting Minutes
Generates professional .docx files from meeting transcripts and summaries
"""

import os
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Dict, Any
import re

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError as e:
    raise ImportError(
        "python-docx is required for Word document generation. "
        "Install it with: pip install python-docx"
    ) from e

from config import logger, _BASE_DIR
from models import Meeting, TranscriptSegment, MeetingSummary
from exceptions import DocumentGenerationError


class WordGenerator:
    """
    Professional Word document generator for meeting minutes
    Supports Korean fonts, custom templates, and structured sections
    """

    # AI Classification Categories (user can select 0-3)
    CATEGORIES = ["현황", "배경", "논의", "문제점", "의견", "결의"]

    def __init__(
        self,
        output_dir: Optional[Path] = None,
        default_font: str = "맑은 고딕",
        default_font_size: int = 11,
        title_font_size: int = 18,
        heading_font_size: int = 14
    ):
        """
        Initialize Word document generator

        Args:
            output_dir: Output directory for generated documents (default: ./output)
            default_font: Default font name for Korean support
            default_font_size: Default font size in points
            title_font_size: Title font size in points
            heading_font_size: Heading font size in points
        """
        self.output_dir = output_dir or (_BASE_DIR / "output")
        self.output_dir = Path(self.output_dir).resolve()
        self.output_dir.mkdir(parents=True, exist_ok=True)

        self.default_font = default_font
        self.default_font_size = default_font_size
        self.title_font_size = title_font_size
        self.heading_font_size = heading_font_size

        logger.info(f"WordGenerator initialized. Output dir: {self.output_dir}")

    def _sanitize_filename(self, filename: str) -> str:
        """
        Sanitize filename by removing invalid characters

        Args:
            filename: Original filename

        Returns:
            Sanitized filename safe for filesystem
        """
        # Remove invalid Windows filename characters
        invalid_chars = r'[<>:"/\\|?*]'
        sanitized = re.sub(invalid_chars, '_', filename)

        # Limit length to avoid path issues
        max_length = 200
        if len(sanitized) > max_length:
            sanitized = sanitized[:max_length]

        return sanitized.strip()

    def _generate_filename(self, meeting: Meeting) -> str:
        """
        Generate filename from meeting metadata

        Args:
            meeting: Meeting object

        Returns:
            Filename in format: 회의록_{title}_{date}.docx
        """
        date_str = meeting.created_at.strftime("%Y%m%d")
        title_clean = self._sanitize_filename(meeting.title)

        filename = f"회의록_{title_clean}_{date_str}.docx"
        return filename

    def _set_run_font(self, run, font_name: str, font_size: int, bold: bool = False, color: Optional[RGBColor] = None):
        """
        Set font properties for a text run

        Args:
            run: Document run object
            font_name: Font name
            font_size: Font size in points
            bold: Whether text should be bold
            color: Optional RGB color
        """
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color

    def _add_title(self, doc: Document, title: str):
        """
        Add document title with formatting

        Args:
            doc: Document object
            title: Title text
        """
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in title_para.runs:
            self._set_run_font(
                run,
                self.default_font,
                self.title_font_size,
                bold=True
            )

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """
        Add section heading with consistent formatting

        Args:
            doc: Document object
            text: Heading text
            level: Heading level (1 or 2)
        """
        heading = doc.add_heading(text, level=level)

        for run in heading.runs:
            self._set_run_font(
                run,
                self.default_font,
                self.heading_font_size if level == 1 else self.default_font_size + 1,
                bold=True,
                color=RGBColor(0, 51, 102)  # Dark blue
            )

    def _add_paragraph(self, doc: Document, text: str, bold: bool = False, indent: float = 0.0):
        """
        Add paragraph with consistent formatting

        Args:
            doc: Document object
            text: Paragraph text
            bold: Whether text should be bold
            indent: Left indent in inches
        """
        para = doc.add_paragraph()
        if indent > 0:
            para.paragraph_format.left_indent = Inches(indent)

        run = para.add_run(text)
        self._set_run_font(run, self.default_font, self.default_font_size, bold=bold)

        return para

    def _add_metadata_section(self, doc: Document, meeting: Meeting, participants: List[str]):
        """
        Add meeting metadata section (date, time, participants)

        Args:
            doc: Document object
            meeting: Meeting object
            participants: List of participant names
        """
        # Date and time
        date_str = meeting.created_at.strftime("%Y년 %m월 %d일")
        time_str = meeting.created_at.strftime("%H:%M")

        self._add_paragraph(doc, f"일시: {date_str} {time_str}", bold=True)

        # Duration if available
        if meeting.duration_seconds:
            duration_minutes = int(meeting.duration_seconds / 60)
            self._add_paragraph(doc, f"소요 시간: {duration_minutes}분", bold=True)

        # Participants
        if participants:
            participants_str = ", ".join(participants)
            self._add_paragraph(doc, f"참석자: {participants_str}", bold=True)

        # Add spacing
        doc.add_paragraph()

    def _add_category_sections(
        self,
        doc: Document,
        summary: MeetingSummary,
        selected_categories: Optional[List[str]] = None
    ):
        """
        Add AI-classified category sections (0-3 categories)

        Args:
            doc: Document object
            summary: MeetingSummary object
            selected_categories: List of category names to include (subset of CATEGORIES)
        """
        if not selected_categories:
            return

        # Validate categories
        valid_categories = [cat for cat in selected_categories if cat in self.CATEGORIES]
        if not valid_categories:
            logger.warning(f"No valid categories selected from: {selected_categories}")
            return

        if len(valid_categories) > 3:
            logger.warning(f"More than 3 categories selected, limiting to first 3")
            valid_categories = valid_categories[:3]

        self._add_heading(doc, "AI 분류", level=1)

        # Map category names to content extraction logic
        category_content_map = {
            "현황": self._extract_status_from_summary,
            "배경": self._extract_background_from_summary,
            "논의": self._extract_discussion_from_summary,
            "문제점": self._extract_issues_from_summary,
            "의견": self._extract_opinions_from_summary,
            "결의": self._extract_resolutions_from_summary,
        }

        for category in valid_categories:
            self._add_heading(doc, category, level=2)

            # Extract content using category-specific logic
            content_extractor = category_content_map.get(category)
            if content_extractor:
                content = content_extractor(summary)
                self._add_paragraph(doc, content)
            else:
                self._add_paragraph(doc, f"({category} 내용)")

            doc.add_paragraph()

    def _extract_status_from_summary(self, summary: MeetingSummary) -> str:
        """Extract status/current situation from summary"""
        # Try to extract status-related content from summary
        # For MVP, use key points or summary itself
        if summary.key_points:
            return "\n".join(f"- {point}" for point in summary.key_points[:2])
        return summary.summary[:200] + "..." if len(summary.summary) > 200 else summary.summary

    def _extract_background_from_summary(self, summary: MeetingSummary) -> str:
        """Extract background/context from summary"""
        # Use summary text as background context
        return summary.summary[:300] + "..." if len(summary.summary) > 300 else summary.summary

    def _extract_discussion_from_summary(self, summary: MeetingSummary) -> str:
        """Extract discussion points from summary"""
        if summary.topics:
            return "\n".join(f"- {topic}" for topic in summary.topics)
        elif summary.key_points:
            return "\n".join(f"- {point}" for point in summary.key_points)
        return summary.summary

    def _extract_issues_from_summary(self, summary: MeetingSummary) -> str:
        """Extract issues/problems from summary"""
        # Look for problem-related key points
        if summary.key_points:
            problem_keywords = ["문제", "이슈", "과제", "리스크"]
            problem_points = [
                point for point in summary.key_points
                if any(keyword in point for keyword in problem_keywords)
            ]
            if problem_points:
                return "\n".join(f"- {point}" for point in problem_points)
        return "(논의된 문제점 및 이슈)"

    def _extract_opinions_from_summary(self, summary: MeetingSummary) -> str:
        """Extract opinions/feedback from summary"""
        # Use key points as opinions
        if summary.key_points:
            return "\n".join(f"- {point}" for point in summary.key_points)
        return "(회의 참석자 의견)"

    def _extract_resolutions_from_summary(self, summary: MeetingSummary) -> str:
        """Extract resolutions/decisions from summary"""
        # Action items are typically resolutions
        if summary.action_items:
            return "\n".join(f"- {item}" for item in summary.action_items)
        return "(결정 사항 및 결의)"

    def _add_transcript_section(self, doc: Document, transcripts: List[TranscriptSegment]):
        """
        Add full conversation transcript with speakers and timestamps

        Args:
            doc: Document object
            transcripts: List of transcript segments
        """
        self._add_heading(doc, "전체 대화 내용", level=1)

        for segment in transcripts:
            # Format speaker and timestamp
            speaker = segment.speaker_label or f"Speaker {segment.speaker_id}" if segment.speaker_id else "Unknown"
            timestamp = self._format_timestamp(segment.start_time)

            # Add speaker line with timestamp
            speaker_para = doc.add_paragraph()
            speaker_run = speaker_para.add_run(f"[{timestamp}] {speaker}: ")
            self._set_run_font(speaker_run, self.default_font, self.default_font_size, bold=True)

            # Add transcript text
            text_run = speaker_para.add_run(segment.text)
            self._set_run_font(text_run, self.default_font, self.default_font_size)

    def _format_timestamp(self, seconds: float) -> str:
        """
        Format seconds to HH:MM:SS timestamp

        Args:
            seconds: Time in seconds

        Returns:
            Formatted timestamp string
        """
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"

    def _add_action_items_section(self, doc: Document, summary: MeetingSummary):
        """
        Add action items section

        Args:
            doc: Document object
            summary: MeetingSummary object with action items
        """
        if not summary.action_items:
            return

        self._add_heading(doc, "액션 아이템", level=1)

        for item in summary.action_items:
            self._add_paragraph(doc, f"• {item}", indent=0.25)

        doc.add_paragraph()

    def _extract_participants_from_transcripts(self, transcripts: List[TranscriptSegment]) -> List[str]:
        """
        Extract unique speakers from transcript segments

        Args:
            transcripts: List of transcript segments

        Returns:
            List of unique speaker names
        """
        speakers = set()
        for segment in transcripts:
            if segment.speaker_label:
                speakers.add(segment.speaker_label)
            elif segment.speaker_id:
                speakers.add(f"Speaker {segment.speaker_id}")

        return sorted(list(speakers))

    def generate_meeting_minutes(
        self,
        meeting: Meeting,
        transcripts: List[TranscriptSegment],
        summary: MeetingSummary,
        selected_categories: Optional[List[str]] = None,
        custom_filename: Optional[str] = None
    ) -> Path:
        """
        Generate complete meeting minutes Word document

        Args:
            meeting: Meeting object with metadata
            transcripts: List of transcript segments
            summary: AI-generated meeting summary
            selected_categories: Optional list of 0-3 categories from CATEGORIES
            custom_filename: Optional custom filename (without extension)

        Returns:
            Path to generated .docx file

        Raises:
            DocumentGenerationError: If document generation fails
        """
        try:
            logger.info(f"Generating meeting minutes for: {meeting.title}")

            # Create new document
            doc = Document()

            # Set default styles
            styles = doc.styles
            normal_style = styles['Normal']
            font = normal_style.font
            font.name = self.default_font
            font.size = Pt(self.default_font_size)

            # 1. Add title
            self._add_title(doc, meeting.title)

            # 2. Add metadata section
            participants = self._extract_participants_from_transcripts(transcripts)
            self._add_metadata_section(doc, meeting, participants)

            # 3. Add summary section
            self._add_heading(doc, "요약", level=1)
            self._add_paragraph(doc, summary.summary)
            doc.add_paragraph()

            # 4. Add key points if available
            if summary.key_points:
                self._add_heading(doc, "핵심 포인트", level=1)
                for point in summary.key_points:
                    self._add_paragraph(doc, f"• {point}", indent=0.25)
                doc.add_paragraph()

            # 5. Add AI classification sections (0-3 categories)
            if selected_categories:
                self._add_category_sections(doc, summary, selected_categories)

            # 6. Add full transcript
            if transcripts:
                self._add_transcript_section(doc, transcripts)

            # 7. Add action items
            self._add_action_items_section(doc, summary)

            # Generate filename
            if custom_filename:
                filename = self._sanitize_filename(custom_filename) + ".docx"
            else:
                filename = self._generate_filename(meeting)

            # Save document
            output_path = self.output_dir / filename
            doc.save(str(output_path))

            logger.info(f"Meeting minutes generated successfully: {output_path}")
            logger.info(f"File size: {output_path.stat().st_size / 1024:.2f} KB")

            return output_path

        except Exception as e:
            error_msg = f"Failed to generate meeting minutes: {e}"
            logger.error(error_msg)
            raise DocumentGenerationError(error_msg) from e

    def generate_template_document(
        self,
        title: str = "회의록 템플릿",
        include_all_sections: bool = True
    ) -> Path:
        """
        Generate a blank template document for reference

        Args:
            title: Template title
            include_all_sections: Whether to include all possible sections

        Returns:
            Path to generated template file
        """
        try:
            doc = Document()

            # Set default styles
            styles = doc.styles
            normal_style = styles['Normal']
            font = normal_style.font
            font.name = self.default_font
            font.size = Pt(self.default_font_size)

            # Add title
            self._add_title(doc, title)

            # Add metadata placeholders
            self._add_paragraph(doc, "일시: YYYY년 MM월 DD일 HH:MM", bold=True)
            self._add_paragraph(doc, "소요 시간: XX분", bold=True)
            self._add_paragraph(doc, "참석자: (참석자 명단)", bold=True)
            doc.add_paragraph()

            # Add section placeholders
            if include_all_sections:
                self._add_heading(doc, "요약", level=1)
                self._add_paragraph(doc, "(회의 요약 내용)")
                doc.add_paragraph()

                self._add_heading(doc, "핵심 포인트", level=1)
                self._add_paragraph(doc, "• (포인트 1)")
                self._add_paragraph(doc, "• (포인트 2)")
                doc.add_paragraph()

                self._add_heading(doc, "AI 분류", level=1)
                for category in self.CATEGORIES[:3]:
                    self._add_heading(doc, category, level=2)
                    self._add_paragraph(doc, f"({category} 내용)")
                    doc.add_paragraph()

                self._add_heading(doc, "전체 대화 내용", level=1)
                self._add_paragraph(doc, "[00:00:00] 발표자: (대화 내용)")
                doc.add_paragraph()

                self._add_heading(doc, "액션 아이템", level=1)
                self._add_paragraph(doc, "• (액션 아이템 1)")

            # Save template
            filename = f"{self._sanitize_filename(title)}.docx"
            output_path = self.output_dir / filename
            doc.save(str(output_path))

            logger.info(f"Template document generated: {output_path}")
            return output_path

        except Exception as e:
            error_msg = f"Failed to generate template: {e}"
            logger.error(error_msg)
            raise DocumentGenerationError(error_msg) from e


def get_word_generator(
    output_dir: Optional[Path] = None,
    default_font: str = "맑은 고딕",
    default_font_size: int = 11
) -> WordGenerator:
    """
    Factory function to create WordGenerator instance

    Args:
        output_dir: Output directory path
        default_font: Default font for Korean support
        default_font_size: Default font size

    Returns:
        WordGenerator instance
    """
    return WordGenerator(
        output_dir=output_dir,
        default_font=default_font,
        default_font_size=default_font_size
    )
